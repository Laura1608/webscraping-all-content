import requests
from bs4 import BeautifulSoup
from selenium import webdriver
import time
import os.path
from urllib.parse import urlparse
from docx import Document
import os
import pandas as pd

# Create empty list to store the scraped content
data_list = []

# Define path to the Excel-file with URLs
sitemap = "vermeertexas.com_sitemap.xlsx"

# Read URLs from the Excel-file
df = pd.read_excel(sitemap)

# Returns data column as a list
urls = df['URL'].tolist()

# Create new Word document to store the scraped content
doc = Document()

# Add title to the document
doc.add_heading('Text and image URLs from all pages VermeerTexas.com', level=1)

# Loop through each URL and start scraping
for url in urls:

    # Initialize Selenium WebDriver
    driver = webdriver.Chrome()

    # Open the web page
    driver.get(url)

    # Before continuing, wait 3 seconds
    time.sleep(3)

    # Retrieve HTML-content from webpage
    response = BeautifulSoup(driver.page_source, 'html.parser')

    # Only retrieve its text
    plain_text = response.get_text(separator="\n", strip=True)

    # When page contains text, add it to the document
    if plain_text:
        doc.add_heading('Plain text copy', level=2)
        doc.add_paragraph(url)
        doc.add_paragraph(plain_text)

    # Define folder to save images
    folder = 'downloaded_images'

    # Create the folder if it doesn't exist yet
    if not os.path.exists(folder):
        os.makedirs(folder)

    # Function to sanitize file names
    def sanitize_filename(filename):
        # Split the filename into name and extension
        name, ext = os.path.splitext(filename)
        # Replace invalid characters in the name part only
        name = "".join(c if c.isalnum() else "_" for c in name)
        # Reattach the extension
        return name + ext


    # Create empty list for URLs
    img_urls = []

    # Find all 'img' tags and extract their 'src' attributes
    imgs = response.find_all('img')
    if imgs:
        for img in imgs:
            src = img.get('src')
            # Only extract images with a valid URL
            if src and (src.startswith('http://') or src.startswith('https://')):
                if 'www.facebook.com' or 'www.instagram.com' or 'www.youtube.com' or 'www.linkedin.com' not in src:
                    img_urls.append(src)

    # Add image URLs to the document
    doc.add_heading('Image URLs', level=2)
    doc.add_paragraph(url)

    # Use page url to create the folder name
    folder_name = url.replace('.com', '').replace('https://www.', '').replace('http://www.', '').replace('/', '-')
    folder_name = folder_name.strip('_')

    # Define dynamic folder path
    dynamic_folder_path = os.path.join(folder, folder_name)

    # Make sure directory exists
    os.makedirs(dynamic_folder_path, exist_ok=True)

    # Loop through all images on the page
    for img_url in img_urls:
        # Add image URLs to the document
        doc.add_paragraph(img_url)

        # Get image response from website
        img_response = requests.get(img_url)

        # Extract the image file name and sanitize it
        parsed_url = urlparse(img_url)
        img_name = sanitize_filename(os.path.basename(parsed_url.path) or parsed_url.netloc)

        # Create img path to save the image
        img_path = os.path.join(dynamic_folder_path, img_name)

        # Save each image to their folder
        with open(img_path, 'wb') as img_file:
            img_file.write(img_response.content)

    # Add a page break after processing URLs
    doc.add_page_break()

    # Close the Selenium browser
    driver.quit()

# Save the Word-document
doc.save('vermeertexas.com_content_pages.docx')
print('Info saved in doc!')
